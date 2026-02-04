"""
File Parser Module
Extracts text content from various document formats
"""

from typing import Optional
import os
from pathlib import Path


def parse_file(filepath: str) -> Optional[str]:
    """
    Parse a file and extract text content.
    
    Args:
        filepath: Path to the file to parse
        
    Returns:
        Extracted text content or None if parsing failed
    """
    file_ext = Path(filepath).suffix.lower()
    
    try:
        if file_ext == '.txt' or file_ext == '.md':
            return parse_text_file(filepath)
        elif file_ext == '.docx':
            return parse_docx_file(filepath)
        elif file_ext == '.pdf':
            return parse_pdf_file(filepath)
        else:
            print(f"Unsupported file type: {file_ext}")
            return None
    except Exception as e:
        print(f"Error parsing file: {e}")
        return None


def parse_text_file(filepath: str) -> Optional[str]:
    """
    Parse plain text or markdown file.
    
    Args:
        filepath: Path to text file
        
    Returns:
        File content as string
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        print(f"✅ Parsed text file: {len(content)} characters")
        return content
    except Exception as e:
        print(f"Error reading text file: {e}")
        return None


def parse_docx_file(filepath: str) -> Optional[str]:
    """
    Parse Word document (.docx).
    
    Args:
        filepath: Path to .docx file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
        
        doc = Document(filepath)
        
        # Extract all paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
        
        # Combine all text
        content = '\n\n'.join(paragraphs)
        if table_text:
            content += '\n\nTables:\n' + '\n'.join(table_text)
        
        print(f"✅ Parsed .docx file: {len(content)} characters")
        return content
        
    except ImportError:
        print("Error: python-docx library not installed")
        return None
    except Exception as e:
        print(f"Error parsing .docx file: {e}")
        return None


def parse_pdf_file(filepath: str) -> Optional[str]:
    """
    Parse PDF file.
    
    Args:
        filepath: Path to .pdf file
        
    Returns:
        Extracted text content
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(filepath)
        
        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(page_text)
        
        content = '\n\n'.join(text_parts)
        print(f"✅ Parsed .pdf file: {len(reader.pages)} pages, {len(content)} characters")
        return content
        
    except ImportError:
        print("Error: PyPDF2 library not installed")
        return None
    except Exception as e:
        print(f"Error parsing .pdf file: {e}")
        return None
